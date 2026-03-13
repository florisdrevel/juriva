from flask import Flask, request, jsonify, render_template, session, send_file
import anthropic
import pypdf
import zipfile
from xml.etree import ElementTree as ET
import io
import os
import re
import json
import sqlite3
import secrets
import string
import smtplib
try:
    import stripe
    STRIPE_AVAILABLE = True
except ImportError:
    stripe = None
    STRIPE_AVAILABLE = False
    print("[STRIPE] stripe package not installed — payments disabled", flush=True)
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime, timedelta
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY") or "REDACTED-FLASK-SECRET"

# ─────────────────────────────────────────────────────────────
# CONFIGURATION — set via Railway environment variables
# ─────────────────────────────────────────────────────────────
ANTHROPIC_API_KEY   = os.environ.get("ANTHROPIC_API_KEY", "REDACTED-ANTHROPIC-KEY")
STRIPE_SECRET_KEY   = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
GMAIL_FROM          = os.environ.get("GMAIL_FROM", "info.juriva@gmail.com")
GMAIL_APP_PASSWORD  = os.environ.get("GMAIL_APP_PASSWORD", "")
RESEND_API_KEY      = os.environ.get("RESEND_API_KEY", "")

# Stripe Price IDs — fill in after creating products in Stripe dashboard
STRIPE_PRICES = {
    "per_analyse_monthly": os.environ.get("STRIPE_PRICE_PER_ANALYSE", ""),
    "zzp_monthly":         os.environ.get("STRIPE_PRICE_ZZP_MONTHLY", ""),
    "zzp_annual":          os.environ.get("STRIPE_PRICE_ZZP_ANNUAL", ""),
    "pro_monthly":         os.environ.get("STRIPE_PRICE_PRO_MONTHLY", ""),
    "pro_annual":          os.environ.get("STRIPE_PRICE_PRO_ANNUAL", ""),
    "kantoor_monthly":     os.environ.get("STRIPE_PRICE_KANTOOR_MONTHLY", ""),
    "kantoor_annual":      os.environ.get("STRIPE_PRICE_KANTOOR_ANNUAL", ""),
}

if STRIPE_SECRET_KEY and STRIPE_AVAILABLE:
    stripe.api_key = STRIPE_SECRET_KEY

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

TRIAL_DAYS = 14

# ─────────────────────────────────────────────────────────────
# SQLITE — persistent storage for codes & activations
# ─────────────────────────────────────────────────────────────
# Use /tmp as fallback if /app isn't writable (safe for Railway)
_default_db = "/app/juriva.db" if os.path.isdir("/app") else "/tmp/juriva.db"
DB_PATH = os.environ.get("DB_PATH", _default_db)

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS codes (
                code TEXT PRIMARY KEY,
                plan TEXT NOT NULL DEFAULT 'pilot',
                email TEXT,
                activated_at TEXT,
                created_at TEXT NOT NULL,
                stripe_session_id TEXT,
                is_subscription INTEGER DEFAULT 0,
                analyse_count INTEGER DEFAULT 0,
                subscription_end TEXT,
                max_users INTEGER DEFAULT 1
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS payments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                stripe_session_id TEXT UNIQUE,
                customer_email TEXT,
                plan TEXT,
                amount_total INTEGER,
                currency TEXT,
                created_at TEXT,
                code_sent TEXT
            )
        """)
        # Migrate: add new columns if missing
        for col, definition in [
            ('subscription_end', 'TEXT'),
            ('max_users', 'INTEGER DEFAULT 1')
        ]:
            try:
                conn.execute(f"ALTER TABLE codes ADD COLUMN {col} {definition}")
            except Exception:
                pass
        conn.commit()
        # Seed pilot codes if not present
        pilot_codes = [
            ("JURIVA-PILOT-1", "pilot"), ("JURIVA-PILOT-2", "pilot"),
            ("JURIVA-PILOT-3", "pilot"), ("JURIVA-PILOT-4", "pilot"),
            ("JURIVA-PILOT-5", "pilot"), ("JURIVA-PILOT-6", "pilot"),
            ("JURIVA-PILOT-7", "pilot"), ("JURIVA-PILOT-8", "pilot"),
            ("JURIVA-PILOT-9", "pilot"), ("JURIVA-PILOT-10", "pilot"),
        ]
        for code, plan in pilot_codes:
            conn.execute(
                "INSERT OR IGNORE INTO codes (code, plan, created_at) VALUES (?, ?, ?)",
                (code, plan, datetime.now().isoformat())
            )
        # Seed existing firm outreach codes
        firm_codes = [
            "JURIVA-DENTONS-001","JURIVA-CLIFFORDCHANCE-002","JURIVA-ALLENOVERY-003",
            "JURIVA-FRESHFIELDS-004","JURIVA-HOUTHOFF-005","JURIVA-DEBRAUW-006",
            "JURIVA-STIBBE-007","JURIVA-NautaDutilh-008","JURIVA-LOYENSLOEFF-009",
            "JURIVA-AKDLAW-010","JURIVA-BOEKEL-011","JURIVA-PLOUMLAW-012",
            "JURIVA-LEXENCECORP-013","JURIVA-LEXENCELITIG-014",
        ]
        for i, code in enumerate(firm_codes):
            conn.execute(
                "INSERT OR IGNORE INTO codes (code, plan, created_at) VALUES (?, ?, ?)",
                (code, "pilot", datetime.now().isoformat())
            )
        conn.commit()

try:
    init_db()
    print(f"[DB] Initialised at {DB_PATH}", flush=True)
except Exception as e:
    print(f"[DB ERROR] Could not init DB: {e}", flush=True)

print(f"[APP] Flask ready — Stripe: {STRIPE_AVAILABLE}, Gmail: {bool(GMAIL_APP_PASSWORD)}", flush=True)

# ─────────────────────────────────────────────────────────────
# CODE HELPERS
# ─────────────────────────────────────────────────────────────
def generate_access_code(plan: str) -> str:
    """Generate a unique random access code for a paid plan."""
    plan_slug = plan.upper().replace(" ", "-")[:12]
    chars = string.ascii_uppercase + string.digits
    random_part = ''.join(secrets.choice(chars) for _ in range(8))
    return f"JRV-{plan_slug}-{random_part}"

def is_code_valid(code: str):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM codes WHERE code = ?", (code,)).fetchone()
    if not row:
        return False, "invalid"
    plan = row["plan"]
    # Per Analyse — valid until 1 analysis is used
    if plan == 'Per Analyse':
        used = row["analyse_count"] or 0
        if used >= 1:
            return False, "expired"
        return True, "ok"
    # Subscriptions with end date (ZZP/Pro/Kantoor monthly or annual)
    if row["is_subscription"] and row["subscription_end"]:
        end = datetime.fromisoformat(row["subscription_end"])
        if datetime.now() > end:
            return False, "expired"
        return True, "ok"
    # Subscriptions without end date — still valid (Stripe manages renewal)
    if row["is_subscription"]:
        return True, "ok"
    # Pilot / trial codes — 14-day window from activation
    if row["activated_at"]:
        activated_at = datetime.fromisoformat(row["activated_at"])
        if datetime.now() > activated_at + timedelta(days=TRIAL_DAYS):
            return False, "expired"
    return True, "ok"

def days_remaining(code: str) -> int:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM codes WHERE code = ?", (code,)).fetchone()
    if not row:
        return 0
    plan = row["plan"]
    if plan == 'Per Analyse':
        used = row["analyse_count"] or 0
        return max(0, 1 - used)
    if row["is_subscription"] and row["subscription_end"]:
        end = datetime.fromisoformat(row["subscription_end"])
        return max(0, (end - datetime.now()).days)
    if row["is_subscription"]:
        return 999
    if not row["activated_at"]:
        return TRIAL_DAYS
    activated_at = datetime.fromisoformat(row["activated_at"])
    return max(0, (activated_at + timedelta(days=TRIAL_DAYS) - datetime.now()).days)

def activate_code(code: str):
    with get_db() as conn:
        row = conn.execute("SELECT activated_at FROM codes WHERE code = ?", (code,)).fetchone()
        if row and not row["activated_at"]:
            conn.execute(
                "UPDATE codes SET activated_at = ? WHERE code = ?",
                (datetime.now().isoformat(), code)
            )
            conn.commit()

# ─────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────
def extract_text_from_file(file):
    filename = file.filename.lower()
    content = file.read()
    if filename.endswith('.pdf'):
        reader = pypdf.PdfReader(io.BytesIO(content))
        text = "".join(page.extract_text() or '' for page in reader.pages)
    elif filename.endswith('.docx'):
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            xml = z.read('word/document.xml')
        tree = ET.fromstring(xml)
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        text = " ".join(node.text for node in tree.iter(f'{ns}t') if node.text)
    else:
        text = content.decode('utf-8')
    del content
    return text

# ─────────────────────────────────────────
# EMAIL SENDER
# ─────────────────────────────────────────
def send_access_email(to_email: str, code: str, plan: str, lang: str = 'nl'):
    """Send access code email via Resend API."""
    if not RESEND_API_KEY:
        print(f"[EMAIL SKIP] No Resend API key. Code for {to_email}: {code}", flush=True)
        return False
    try:
        import requests as _requests
        if lang == 'nl':
            subject = f"Uw Juriva toegangscode — {plan}"
            html_body = f"""<div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;color:#1a1815">
  <div style="border-bottom:2px solid #c8a951;padding-bottom:12px;margin-bottom:24px">
    <span style="font-size:22px;font-weight:700;letter-spacing:-0.02em">Juriva</span>
  </div>
  <p>Beste,</p>
  <p>Bedankt voor uw aankoop van <strong>Juriva {plan}</strong>.</p>
  <p>Uw toegangscode is:</p>
  <div style="background:#f5f3ef;border:1px solid #e0dbd2;border-radius:8px;padding:16px 24px;margin:20px 0;text-align:center">
    <span style="font-size:22px;font-weight:700;letter-spacing:0.08em;color:#1a1815">{code}</span>
  </div>
  <p><a href="https://juriva.nl" style="background:#c8a951;color:#fff;padding:10px 22px;border-radius:6px;text-decoration:none;font-weight:600">Ga naar juriva.nl →</a></p>
  <p style="margin-top:24px">Heeft u vragen? Stuur een e-mail naar <a href="mailto:info@juriva.nl">info@juriva.nl</a> — wij reageren binnen 1 werkdag.</p>
  <p>Met vriendelijke groet,<br><strong>Florean Drevel</strong><br>Juriva · juriva.nl</p>
  <p style="font-size:11px;color:#999;margin-top:32px;border-top:1px solid #eee;padding-top:12px">Juriva levert geen juridisch advies. Contracten dienen altijd door een bevoegde advocaat te worden beoordeeld.</p>
</div>"""
        else:
            subject = f"Your Juriva access code — {plan}"
            html_body = f"""<div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;color:#1a1815">
  <div style="border-bottom:2px solid #c8a951;padding-bottom:12px;margin-bottom:24px">
    <span style="font-size:22px;font-weight:700;letter-spacing:-0.02em">Juriva</span>
  </div>
  <p>Hello,</p>
  <p>Thank you for purchasing <strong>Juriva {plan}</strong>.</p>
  <p>Your access code is:</p>
  <div style="background:#f5f3ef;border:1px solid #e0dbd2;border-radius:8px;padding:16px 24px;margin:20px 0;text-align:center">
    <span style="font-size:22px;font-weight:700;letter-spacing:0.08em;color:#1a1815">{code}</span>
  </div>
  <p><a href="https://juriva.nl" style="background:#c8a951;color:#fff;padding:10px 22px;border-radius:6px;text-decoration:none;font-weight:600">Go to juriva.nl →</a></p>
  <p style="margin-top:24px">Questions? Email <a href="mailto:info@juriva.nl">info@juriva.nl</a> — we respond within 1 business day.</p>
  <p>Kind regards,<br><strong>Florean Drevel</strong><br>Juriva · juriva.nl</p>
  <p style="font-size:11px;color:#999;margin-top:32px;border-top:1px solid #eee;padding-top:12px">Juriva does not provide legal advice. Contracts should always be reviewed by a qualified lawyer.</p>
</div>"""

        resp = _requests.post(
            "https://api.resend.com/emails",
            json={
                "from": "Juriva <info@juriva.nl>",
                "to": [to_email],
                "subject": subject,
                "html": html_body
            },
            headers={"Authorization": f"Bearer {RESEND_API_KEY}"},
            timeout=15
        )
        print(f"[EMAIL] Resend status: {resp.status_code} — {resp.text}", flush=True)
        resp.raise_for_status()
        print(f"[EMAIL OK] Sent to {to_email}", flush=True)
        return True
    except Exception as e:
        print(f"[EMAIL ERROR] {e}", flush=True)
        import traceback; traceback.print_exc()
        return False

# ─────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────

@app.after_request
def disable_cf_email_obfuscation(response):
    # Allow Google to index public pages, block caching on API responses
    if request.path.startswith('/api') or request.path.startswith('/webhook') or request.path.startswith('/admin'):
        response.headers['X-Robots-Tag'] = 'noindex'
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return response

@app.errorhandler(400)
def bad_request(e): return jsonify({'error': str(e)}), 400
@app.errorhandler(401)
def unauthorized(e): return jsonify({'error': str(e)}), 401
@app.errorhandler(403)
def forbidden(e): return jsonify({'error': str(e)}), 403
@app.errorhandler(404)
def not_found(e): return jsonify({'error': str(e)}), 404
@app.errorhandler(500)
def server_error(e): return jsonify({'error': str(e)}), 500
@app.errorhandler(Exception)
def unhandled(e): return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/sitemap.xml')
def sitemap():
    from flask import Response
    xml = open(os.path.join(os.path.dirname(__file__), 'sitemap.xml')).read()
    return Response(xml, mimetype='application/xml')

@app.route('/robots.txt')
def robots():
    from flask import Response
    txt = open(os.path.join(os.path.dirname(__file__), 'robots.txt')).read()
    return Response(txt, mimetype='text/plain')

@app.route('/api/verify-code', methods=['POST'])
def verify_code():
    data = request.get_json()
    code = data.get('code', '').strip().upper()
    valid, reason = is_code_valid(code)
    if reason == "expired":
        return jsonify({'success': False, 'error': 'expired'})
    if not valid:
        return jsonify({'success': False, 'error': 'invalid'})
    activate_code(code)
    # Get plan from DB and store in session
    with get_db() as conn:
        row = conn.execute("SELECT plan FROM codes WHERE code = ?", (code,)).fetchone()
        plan = row['plan'] if row else 'pilot'
    session['authenticated'] = True
    session['code'] = code
    session['plan'] = plan
    session['terms_accepted'] = session.get('terms_accepted', False)
    # Per Analyse: track usage count
    if plan == 'Per Analyse':
        with get_db() as conn:
            row = conn.execute("SELECT analyse_count FROM codes WHERE code = ?", (code,)).fetchone()
            used = row['analyse_count'] if row and 'analyse_count' in row.keys() else 0
        session['analyses_remaining'] = max(0, 1 - used)
    else:
        session['analyses_remaining'] = 999
    return jsonify({
        'success': True,
        'days_remaining': days_remaining(code),
        'plan': plan,
        'analyses_remaining': session.get('analyses_remaining', 999)
    })

@app.route('/api/check-auth', methods=['GET'])
def check_auth():
    if not session.get('authenticated'):
        return jsonify({'authenticated': False})
    code = session.get('code', '')
    valid, reason = is_code_valid(code)
    if not valid:
        session.clear()
        return jsonify({'authenticated': False, 'reason': reason})
    # Get analyses remaining for Per Analyse plan
    plan = session.get('plan', 'pilot')
    analyses_remaining = 999
    if plan == 'Per Analyse':
        with get_db() as conn:
            row = conn.execute("SELECT analyse_count FROM codes WHERE code = ?", (code,)).fetchone()
            used = 0
            try: used = row['analyse_count'] if row else 0
            except Exception: used = 0
        analyses_remaining = max(0, 1 - used)

    return jsonify({
        'authenticated': True,
        'terms_accepted': session.get('terms_accepted', False),
        'days_remaining': days_remaining(code),
        'plan': plan,
        'analyses_remaining': analyses_remaining
    })

@app.route('/api/accept-terms', methods=['POST'])
def accept_terms():
    if not session.get('authenticated'):
        return jsonify({'error': 'Not authenticated'}), 401
    session['terms_accepted'] = True
    return jsonify({'success': True})

def purge_session_data():
    for key in ['last_document_text', 'last_playbook_text', 'document_chunks']:
        session.pop(key, None)

@app.route('/api/review', methods=['POST'])
def review_contract():
    if not session.get('authenticated'):
        return jsonify({'error': 'Geen toegang.'}), 401
    if not session.get('terms_accepted'):
        return jsonify({'error': 'terms_not_accepted'}), 403
    code = session.get('code', '')
    valid, reason = is_code_valid(code)
    if not valid:
        session.clear()
        return jsonify({'error': 'expired' if reason == 'expired' else 'Geen toegang.'}), 401

    # Plan enforcement
    plan = session.get('plan', 'pilot')
    is_per_analyse = (plan == 'Per Analyse')
    is_pro_or_above = plan in ('Professioneel', 'Kantoor')

    # Per Analyse: check usage limit
    if is_per_analyse:
        with get_db() as conn:
            row = conn.execute("SELECT analyse_count FROM codes WHERE code = ?", (code,)).fetchone()
            used = 0
            try:
                used = row['analyse_count'] if row else 0
            except Exception:
                used = 0
        if used >= 1:
            return jsonify({'error': 'limit_reached', 'message': 'U heeft uw analyse gebruikt. Koop een nieuw Per Analyse of upgrade naar een abonnement.'}), 403

    try:
        if 'contract' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        file = request.files['contract']
        if not file.filename:
            return jsonify({'error': 'No file selected'}), 400

        contract_text = extract_text_from_file(file)
        if not contract_text or len(contract_text.strip()) < 50:
            return jsonify({'error': 'Could not extract text from file'}), 400

        prompt = request.form.get('prompt', '')

        playbook_text = None
        playbook_context = ""
        if 'playbook' in request.files and request.files['playbook'].filename:
            if not is_pro_or_above:
                return jsonify({'error': 'upgrade_required', 'message': 'Playbook vergelijking is beschikbaar vanaf het Professioneel plan.'}), 403
            playbook_text = extract_text_from_file(request.files['playbook'])
            playbook_context = f"""
PLAYBOOK INSTRUCTIONS:
The user has uploaded their firm's Gold Standard playbook.
Treat deviations from this playbook as the highest priority risks — above all other general risks.
Flag every clause that contradicts or deviates from the playbook standards.
Quote both the playbook standard AND the contract clause.

PLAYBOOK CONTENT:
{playbook_text[:3000]}

"""

        second_doc_text = None
        cross_ref_context = ""
        if 'second_document' in request.files and request.files['second_document'].filename:
            if not is_pro_or_above:
                return jsonify({'error': 'upgrade_required', 'message': 'Multi-document analyse is beschikbaar vanaf het Professioneel plan.'}), 403
            second_doc_text = extract_text_from_file(request.files['second_document'])
            cross_ref_context = f"""
CROSS-REFERENCE INSTRUCTIONS:
The user has uploaded a second document (e.g. SOW, appendix, or side agreement).
Compare both documents and flag every conflict including:
- Conflicting payment terms or amounts
- Conflicting dates or deadlines
- Conflicting liability caps
- Conflicting defined terms
- Anything promised in one document but missing in the other

SECOND DOCUMENT:
{second_doc_text[:3000]}

Add a section called ## DOCUMENT CONFLICTS at the end of your analysis.
"""

        full_prompt = f"{playbook_context}{cross_ref_context}{prompt}"

        system_prompt = """You are a senior legal contract reviewer with 20 years of experience, writing for an audience of lawyers and legal professionals.

Your readers already know the fundamentals of contract law. Never state the obvious.

FILTER — SKIP ANYTHING THAT IS:
- True of every contract (e.g. "parties must comply", "breach may give rise to damages")
- A general description of what a clause does without comparing it to market standard
- Advice a first-year law student would give
- Padding to make the report look thorough

ONLY FLAG:
- Clauses that deviate from Dutch/EU market standard for this specific contract type
- Asymmetries that are unusual for this type of deal
- Missing clauses that are specifically expected in this contract type
- Concrete financial, operational or IP risks with real-world consequences
Always state: what is the market standard, and exactly how does this contract deviate?

SCORING CALIBRATION:
1 — Fully standard. Simple NDA, no asymmetry.
2 — Market-conforming. One minor deviation.
3 — Routine commercial contract. A few non-standard points.
4 — Several non-standard clauses. Some asymmetry.
5 — Multiple problematic clauses. Real exposure in one area.
6 — Clear imbalance. Material financial or operational risk.
7 — Serious risks across multiple areas. Compounding problems.
8 — High danger. Multiple compounding risks.
9 — Extremely dangerous. Existential clauses.
10 — Predatory. Designed to trap the signing party.

Only assign 7+ if you can name TWO clauses each independently justifying it.
Only assign 8+ if those clauses compound each other.
Most commercial contracts score 3-5. Anchor there first.

PRIVACY: You process documents in memory only. Never reference storing or saving documents.
You do not hallucinate. If a contract is clean, say so.
CRITICAL: Stop immediately after the ONTBREKENDE STANDAARDCLAUSULES / MISSING STANDARD CLAUSES section. Do not add conclusions, warnings, disclaimers or any other content."""

        # Capture for cleanup after streaming
        _contract = contract_text
        _playbook = playbook_text
        _second  = second_doc_text

        _code = session.get('code', '')
        _plan = session.get('plan', '')

        def generate():
            try:
                with client.messages.stream(
                    model="claude-haiku-4-5-20251001",
                    max_tokens=8192,
                    temperature=0.1,
                    system=system_prompt,
                    messages=[{"role": "user", "content": f"{full_prompt}\n\nCONTRACT TEXT:\n{_contract}"}]
                ) as stream:
                    for text in stream.text_stream:
                        yield f"data: {json.dumps({'token': text})}\n\n"
                # Increment analyse_count for Per Analyse plan
                if _plan == 'Per Analyse' and _code:
                    with get_db() as conn:
                        conn.execute(
                            "UPDATE codes SET analyse_count = analyse_count + 1 WHERE code = ?",
                            (_code,)
                        )
                        conn.commit()
                yield "data: [DONE]\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
            finally:
                purge_session_data()

        del contract_text
        if playbook_text: del playbook_text
        if second_doc_text: del second_doc_text

        return app.response_class(generate(), mimetype='text/event-stream',
            headers={'X-Accel-Buffering': 'no', 'Cache-Control': 'no-cache'})

    except Exception as e:
        purge_session_data()
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────
# DOCX REPORT DOWNLOAD
# ─────────────────────────────────────────
def _add_border_bottom(para, color='C8B89A', size=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)

def _add_border_left(para, color='C8B89A', size=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    l = OxmlElement('w:left')
    l.set(qn('w:val'), 'single'); l.set(qn('w:sz'), str(size))
    l.set(qn('w:space'), '8'); l.set(qn('w:color'), color)
    pBdr.append(l); pPr.append(pBdr)

def _add_runs(para, text, font='Arial', size=10, bold=False, italic=False, color=None):
    color = color or RGBColor(0x3A, 0x34, 0x2C)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part: continue
        r = para.add_run(part)
        r.font.name = font; r.font.size = Pt(size)
        r.font.bold = bold or (i % 2 == 1)
        r.font.italic = italic; r.font.color.rgb = color

def build_docx_report(analysis_text, lang='nl', contract_filename=''):
    ACCENT = RGBColor(0xC8, 0xB8, 0x9A)
    DARK   = RGBColor(0x1A, 0x18, 0x15)
    BODY   = RGBColor(0x3A, 0x34, 0x2C)
    MUTED  = RGBColor(0x8A, 0x7E, 0x70)
    DIMMED = RGBColor(0xB0, 0xA4, 0x94)
    QUOTE  = RGBColor(0x7A, 0x6C, 0x5A)

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # Header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Juriva')
    r.font.name = 'Georgia'; r.font.size = Pt(30); r.font.color.rgb = DARK

    tagline = 'Een tweede paar ogen dat nooit moe wordt.' if lang == 'nl' else 'A second pair of eyes that never gets tired.'
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(tagline)
    r.font.name = 'Georgia'; r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = MUTED

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    _add_border_bottom(p, 'C8B89A', 6)

    # Meta
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    label = 'CONTRACTANALYSE RAPPORT' if lang == 'nl' else 'CONTRACT ANALYSIS REPORT'
    r = p.add_run(label + '   ')
    r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = ACCENT
    r = p.add_run(datetime.now().strftime('%d %B %Y'))
    r.font.name = 'Arial'; r.font.size = Pt(8); r.font.color.rgb = DIMMED

    if contract_filename:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(20)
        lbl = 'Bestand: ' if lang == 'nl' else 'File: '
        r = p.add_run(lbl); r.font.name = 'Arial'; r.font.size = Pt(8.5); r.font.color.rgb = DIMMED
        r = p.add_run(contract_filename); r.font.name = 'Arial'; r.font.size = Pt(8.5); r.font.color.rgb = MUTED

    # Sections
    for section in [s for s in re.split(r'\n##\s+', '\n' + analysis_text.strip()) if s.strip()]:
        lines = section.strip().split('\n')
        h = doc.add_paragraph(); h.paragraph_format.space_before = Pt(18); h.paragraph_format.space_after = Pt(8)
        _add_border_bottom(h, 'C8B89A', 4)
        r = h.add_run(lines[0].strip().upper())
        r.font.name = 'Arial'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = ACCENT

        for line in lines[1:]:
            line = line.strip()
            if not line: continue
            if re.match(r'^- (CITAAT|QUOTE):', line):
                text = re.sub(r'^- (CITAAT|QUOTE):\s*"?', '', line).rstrip('"')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
                _add_border_left(p, 'C8B89A', 8)
                r = p.add_run('\u201c' + text + '\u201d')
                r.font.name = 'Georgia'; r.font.size = Pt(9.5); r.font.italic = True; r.font.color.rgb = QUOTE
            elif m := re.match(r'^- (ARTIKEL|RISICO|IMPACT|TOELICHTING|ARTICLE|RISK|EXPLANATION|IMPACT|MISSING):\s*(.*)', line):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(m.group(1) + ': ')
                r.font.name = 'Arial'; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = DARK
                r = p.add_run(m.group(2)); r.font.name = 'Arial'; r.font.size = Pt(9.5); r.font.color.rgb = BODY
            elif re.match(r'^\d+\.', line):
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
                _add_runs(p, re.sub(r'^\d+\.\s*', '', line), color=BODY)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
                _add_runs(p, line[2:], color=BODY)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(5)
                _add_runs(p, line, color=BODY)

    # Footer
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)
    _add_border_bottom(p, 'C8B89A', 4)
    disc = ('Dit rapport vormt geen juridisch advies. Altijd laten beoordelen door een bevoegde advocaat.'
            if lang == 'nl' else
            'This report does not constitute legal advice. Always have contracts reviewed by a qualified lawyer.')
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(disc); r.font.name = 'Arial'; r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = DIMMED
    p = doc.add_paragraph()
    r = p.add_run('juriva.nl  \u00b7  info.juriva@gmail.com')
    r.font.name = 'Arial'; r.font.size = Pt(7.5); r.font.color.rgb = ACCENT

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

@app.route('/api/download-report', methods=['POST'])
def download_report():
    if not session.get('authenticated'):
        return jsonify({'error': 'Geen toegang.'}), 401
    if not session.get('terms_accepted'):
        return jsonify({'error': 'terms_not_accepted'}), 403
    try:
        data = request.get_json()
        analysis_text = data.get('analysis', '')
        lang = data.get('lang', 'nl')
        contract_filename = data.get('filename', '')
        if not analysis_text:
            return jsonify({'error': 'No analysis provided'}), 400
        buf = build_docx_report(analysis_text, lang=lang, contract_filename=contract_filename)
        safe = re.sub(r'[^a-zA-Z0-9_-]', '_', contract_filename.replace(' ', '_'))[:40] if contract_filename else 'rapport'
        download_name = f"Juriva_{safe}_{datetime.now().strftime('%Y-%m-%d')}.docx"
        return send_file(buf, as_attachment=True, download_name=download_name,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────
# DEMO ROUTE — no auth required, preloaded NexusCloud contract
# ─────────────────────────────────────────

DEMO_CONTRACT = """
NEXUSCLOUD PLATFORM SERVICES AGREEMENT
Tussen: NexusCloud B.V. ("Leverancier") en Klant B.V. ("Klant")

Artikel 1 – Definities
1.1 "Platform" betekent de door Leverancier aangeboden SaaS-diensten.
1.2 "Klantdata" betekent alle data die Klant uploadt of genereert via het Platform.
1.3 "Beheerdata" betekent data die Leverancier naar eigen inzicht kwalificeert als operationeel relevant.

Artikel 2 – Dienstverlening
2.1 Leverancier verleent Klant een niet-exclusieve, niet-overdraagbare licentie tot gebruik van het Platform.
2.2 Leverancier behoudt het recht functionaliteit te wijzigen zonder compensatie aan Klant.
2.3 Leidende Partij bij interpretatie van deze overeenkomst is uitsluitend de Leverancier.

Artikel 3 – Vergoeding
3.1 Klant betaalt EUR 18.500 per maand.
3.2 Klant heeft geen recht op ontbinding wegens prijswijziging.
3.6 Leverancier past jaarlijks een "Annual Business Value Adjustment" toe van maximaal 25%, gebaseerd op de door Leverancier vastgestelde economische waarde die Klant aan het Platform ontleent. Deze aanpassing vereist geen nadere onderbouwing.

Artikel 4 – Intellectueel Eigendom
4.1 Alle Afgeleide Werken en analyses op basis van Klantdata worden eigendom van Leverancier.
4.2 Klant verleent Leverancier een onherroepelijke, eeuwigdurende, wereldwijde, royaltyvrije, sublicentieerbare en overdraagbare licentie om de Klantdata te gebruiken voor (i) verbetering van het Platform; (ii) ontwikkeling van nieuwe producten en diensten; (iii) commerciële doeleinden inclusief doorverkoop van geaggregeerde datasets aan derden; en (iv) training van kunstmatige-intelligentiemodellen. Deze licentie blijft van kracht na beëindiging.
4.5 Bedrijfsspecifieke inzichten gegenereerd uit Klantdata blijven eigendom van Leverancier.

Artikel 5 – Gegevensverwerking
5.2 Leverancier treedt op als verwerker doch behoudt het recht zelfstandig als verwerkingsverantwoordelijke op te treden voor Klantdata die Leverancier kwalificeert als Beheerdata. De kwalificatie van data als Beheerdata wordt uitsluitend door Leverancier bepaald.
5.4 Na beëindiging vervalt het recht van Klant op toegang tot de Klantdata binnen 48 uur na de beëindigingsdatum. Leverancier is niet verplicht Klantdata te retourneren in een gestructureerd, veelgebruikt of machineleesbaar formaat. Klant heeft geen recht op een exportperiode tenzij Klant een aanvullende Data Retrieval Service afneemt voor EUR 12.500 per 30 dagen.

Artikel 7 – Beëindiging
7.2 Leverancier kan de overeenkomst beëindigen wegens (vi) strategische noodzaak, zonder nadere toelichting.
7.4 Bij vroegtijdige beëindiging door Klant, ook bij toerekenbare tekortkoming van Leverancier, is Klant een vergoeding verschuldigd gelijk aan de resterende looptijd.

Artikel 8 – Vertrouwelijkheid
8.4 Klant erkent dat een schending van artikel 8.1 een onherstelbare schade voor Leverancier oplevert. Bij elke gestelde of dreigende schending kan Leverancier aanspraak maken op een contractuele boete van EUR 500.000 per incident.

Artikel 9 – Aansprakelijkheid
9.2 Leverancier geeft geen garanties met betrekking tot beschikbaarheid, juistheid of continuïteit van het Platform.

Artikel 10 – Toepasselijk Recht
10.1 Op deze overeenkomst is het recht van de staat Delaware (USA) van toepassing.
10.3 Leverancier mag geschillen aanhangig maken bij elke bevoegde rechter wereldwijd. Klant is gebonden aan arbitrage in Wilmington, Delaware.

Artikel 11 – Overige Bepalingen
11.1 Leverancier kan de overeenkomst op elk moment eenzijdig wijzigen door Klant een kennisgeving te sturen met een termijn van 3 kalenderdagen. Voortgezet gebruik geldt als onherroepelijke aanvaarding.
11.2 Klant is een retentievergoeding verschuldigd van 20% van de jaarlijkse vergoeding per jaar dat een Concurrerende Relatie voortduurt, naar het oordeel van Leverancier.
11.3 Bij indienstneming van een medewerker van Leverancier is Klant een boete verschuldigd van EUR 250.000 per persoon.
"""

DEMO_SYSTEM_PROMPT = """You are a senior legal contract reviewer with 20 years of experience, writing for an audience of lawyers and legal professionals.

Your readers already know the fundamentals of contract law. Never state the obvious.

FILTER — SKIP ANYTHING THAT IS:
- True of every contract
- A general description of what a clause does without comparing it to market standard
- Advice a first-year law student would give

ONLY FLAG:
- Clauses that deviate from Dutch/EU market standard for this specific contract type
- Asymmetries that are unusual for this type of deal
- Missing clauses specifically expected in this contract type
- Concrete financial, operational or IP risks with real-world consequences

SCORING CALIBRATION:
1=Standard NDA. 3=Routine commercial. 5=Real exposure one area. 7=Serious risks multiple areas. 9=Existential clauses. 10=Predatory.
Only assign 7+ if you can name TWO clauses each independently justifying it.

CRITICAL: Stop immediately after the ONTBREKENDE STANDAARDCLAUSULES section. Do not add conclusions, warnings, disclaimers or any other content."""

DEMO_PROMPT_NL = """Analyseer dit contract in het Nederlands. Gebruik exact dit formaat:

## RISICOSCORE
SCORE:[cijfer]
[één zin samenvatting van de twee grootste risico's]

---

## TOP 3 RISICO'S

RISICO 1: [titel, max 5 woorden]
CITAAT: [max 20 woorden uit contract]
ARTIKEL: [nummer]
MARKTSTANDAARD: [max 12 woorden]
AFWIJKING: [max 15 woorden]
IMPACT: [max 15 woorden]

RISICO 2: [titel, max 5 woorden]
CITAAT: [max 20 woorden]
ARTIKEL: [nummer]
MARKTSTANDAARD: [max 12 woorden]
AFWIJKING: [max 15 woorden]
IMPACT: [max 15 woorden]

RISICO 3: [titel, max 5 woorden]
CITAAT: [max 20 woorden]
ARTIKEL: [nummer]
MARKTSTANDAARD: [max 12 woorden]
AFWIJKING: [max 15 woorden]
IMPACT: [max 15 woorden]

---

## SAMENVATTING
[Max 3 punten. Elk max 15 woorden. NOOIT herhalen wat al in TOP 3 staat.]
1. [aanvullende bevinding]
2. [aanvullende bevinding]
3. [aanvullende bevinding]

---

## ONGEBRUIKELIJKE CLAUSULES

CLAUSULE 1: [titel, max 5 woorden]
CITAAT: [max 15 woorden]
ARTIKEL: [nummer]
TOELICHTING: [max 15 woorden]

CLAUSULE 2: [titel, max 5 woorden]
CITAAT: [max 15 woorden]
ARTIKEL: [nummer]
TOELICHTING: [max 15 woorden]

CLAUSULE 3: [titel, max 5 woorden]
CITAAT: [max 15 woorden]
ARTIKEL: [nummer]
TOELICHTING: [max 15 woorden]

---

## AANBEVOLEN ACTIES
[Max 5 acties. Elk max 15 woorden.]
1. [actie]
2. [actie]
3. [actie]
4. [actie]
5. [actie]

---

## ONTBREKENDE STANDAARDCLAUSULES
[Max 3 items. Elk max 20 woorden.]
1. [ontbrekende clausule]
2. [ontbrekende clausule]
3. [ontbrekende clausule]"""

DEMO_PROMPT_EN = """Analyse this contract in English. Use exactly this format:

## RISK SCORE
SCORE:[number]
[one sentence summary of the two biggest risks]

---

## TOP 3 RISKS

RISK 1: [title, max 5 words]
QUOTE: [max 20 words from contract]
ARTICLE: [number]
MARKET STANDARD: [max 12 words]
DEVIATION: [max 15 words]
IMPACT: [max 15 words]

RISK 2: [title, max 5 words]
QUOTE: [max 20 words]
ARTICLE: [number]
MARKET STANDARD: [max 12 words]
DEVIATION: [max 15 words]
IMPACT: [max 15 words]

RISK 3: [title, max 5 words]
QUOTE: [max 20 words]
ARTICLE: [number]
MARKET STANDARD: [max 12 words]
DEVIATION: [max 15 words]
IMPACT: [max 15 words]

---

## SUMMARY
[Max 3 points. Each max 15 words. NEVER repeat what is in TOP 3.]
1. [additional finding]
2. [additional finding]
3. [additional finding]

---

## UNUSUAL CLAUSES

CLAUSE 1: [title, max 5 words]
QUOTE: [max 15 words]
ARTICLE: [number]
EXPLANATION: [max 15 words]

CLAUSE 2: [title, max 5 words]
QUOTE: [max 15 words]
ARTICLE: [number]
EXPLANATION: [max 15 words]

CLAUSE 3: [title, max 5 words]
QUOTE: [max 15 words]
ARTICLE: [number]
EXPLANATION: [max 15 words]

---

## RECOMMENDED ACTIONS
[Max 5 actions. Each max 15 words.]
1. [action]
2. [action]
3. [action]
4. [action]
5. [action]

---

## MISSING STANDARD CLAUSES
[Max 3 items. Each max 20 words.]
1. [missing clause]
2. [missing clause]
3. [missing clause]"""

@app.route('/api/demo', methods=['POST'])
def demo_analysis():
    """Public demo endpoint — no auth required. Uses preloaded NexusCloud contract."""
    data = request.get_json() or {}
    lang = data.get('lang', 'nl')
    prompt = DEMO_PROMPT_NL if lang == 'nl' else DEMO_PROMPT_EN

    def generate():
        try:
            with client.messages.stream(
                model="claude-haiku-4-5-20251001",
                max_tokens=4096,
                temperature=0.1,
                system=DEMO_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": f"{prompt}\n\nCONTRACT TEXT:\n{DEMO_CONTRACT}"}]
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'token': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return app.response_class(generate(), mimetype='text/event-stream',
        headers={'X-Accel-Buffering': 'no', 'Cache-Control': 'no-cache'})


@app.route('/pricing')
def pricing():
    return render_template('pricing.html')


# ─────────────────────────────────────────────────────────────
# STRIPE — Checkout session creation
# ─────────────────────────────────────────────────────────────

PLAN_NAMES = {
    "per_analyse_monthly": "Per Analyse",
    "zzp_monthly":         "ZZP",
    "zzp_annual":          "ZZP",
    "pro_monthly":         "Professioneel",
    "pro_annual":          "Professioneel",
    "kantoor_monthly":     "Kantoor",
    "kantoor_annual":      "Kantoor",
}

@app.route('/create-checkout-session', methods=['POST'])
def create_checkout_session():
    if not STRIPE_SECRET_KEY or not STRIPE_AVAILABLE:
        return jsonify({'error': 'Payments not configured yet'}), 503
    data = request.get_json() or {}
    plan_key = data.get('plan', '')
    lang = data.get('lang', 'nl')

    price_id = STRIPE_PRICES.get(plan_key, '')
    if not price_id:
        return jsonify({'error': f'Unknown plan: {plan_key}'}), 400

    is_subscription = 'monthly' in plan_key or 'annual' in plan_key
    # Per analyse is a one-time payment
    mode = 'payment' if plan_key == 'per_analyse_monthly' else 'subscription'

    try:
        base_url = request.host_url.rstrip('/')
        session_obj = stripe.checkout.Session.create(
            line_items=[{'price': price_id, 'quantity': 1}],
            mode=mode,
            success_url=f"{base_url}/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{base_url}/pricing",
            metadata={'plan_key': plan_key, 'lang': lang},
            locale='nl' if lang == 'nl' else 'en',
        )
        return jsonify({'url': session_obj.url})
    except stripe.error.StripeError as e:
        return jsonify({'error': str(e)}), 400


# ─────────────────────────────────────────────────────────────
# STRIPE — Webhook (handles payment completion)
# ─────────────────────────────────────────────────────────────

def get_subscription_end(plan_key: str):
    """Calculate subscription end date based on plan type."""
    if plan_key == 'per_analyse_monthly':
        return None  # no expiry date, controlled by analyse_count
    if 'annual' in plan_key:
        return (datetime.now() + timedelta(days=365)).isoformat()
    if 'monthly' in plan_key:
        return (datetime.now() + timedelta(days=31)).isoformat()
    return None  # pilot codes — no end date

def get_max_users(plan_key: str) -> int:
    """Return max concurrent users for plan."""
    if 'kantoor' in plan_key:
        return 10
    if 'pro' in plan_key:
        return 3
    return 1  # ZZP, Per Analyse, pilot


@app.route('/webhook', methods=['POST'])
def stripe_webhook():
    import json as _json, traceback as _tb
    print("[WEBHOOK] A - entering route", flush=True)
    try:
        payload = request.get_data()
        print(f"[WEBHOOK] B - payload {len(payload)} bytes", flush=True)
        sig_header = request.headers.get('Stripe-Signature', '')
        print(f"[WEBHOOK] C - sig={bool(sig_header)} secret={bool(STRIPE_WEBHOOK_SECRET)}", flush=True)
        try:
            if STRIPE_WEBHOOK_SECRET and sig_header:
                event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
                print("[WEBHOOK] D - signature verified", flush=True)
            else:
                event = _json.loads(payload)
                print("[WEBHOOK] D - parsed without verification", flush=True)
        except Exception as e:
            print(f"[WEBHOOK] FAIL parse: {e}", flush=True)
            return jsonify({'error': str(e)}), 400

        event_type = event.get('type', 'unknown')
        print(f"[WEBHOOK] E - event type: {event_type}", flush=True)

        if event_type == 'checkout.session.completed':
            session_obj = event['data']['object']
            print(f"[WEBHOOK] F - calling handler", flush=True)
            try:
                _handle_successful_payment(session_obj)
                print("[WEBHOOK] G - handler done", flush=True)
            except Exception as e:
                print(f"[WEBHOOK] FAIL handler: {e}", flush=True)
                _tb.print_exc()

        return jsonify({'status': 'ok'})

    except Exception as e:
        print(f"[WEBHOOK] FATAL: {e}", flush=True)
        _tb.print_exc()
        return jsonify({'status': 'ok'})


def _handle_successful_payment(session_obj):
    session_id    = session_obj.get('id', '')
    customer_email = session_obj.get('customer_details', {}).get('email', '')
    plan_key      = session_obj.get('metadata', {}).get('plan_key', 'unknown')
    lang          = session_obj.get('metadata', {}).get('lang', 'nl')
    amount_total  = session_obj.get('amount_total', 0)
    currency      = session_obj.get('currency', 'eur')
    plan_name     = PLAN_NAMES.get(plan_key, plan_key)
    is_subscription = plan_key != 'per_analyse_monthly'

    # Check if already processed (idempotency)
    with get_db() as conn:
        existing = conn.execute(
            "SELECT id FROM payments WHERE stripe_session_id = ?", (session_id,)
        ).fetchone()
        if existing:
            print(f"[WEBHOOK] Already processed session {session_id}", flush=True)
            return

    # Generate unique code
    code = generate_access_code(plan_key)
    while True:
        with get_db() as conn:
            exists = conn.execute("SELECT 1 FROM codes WHERE code = ?", (code,)).fetchone()
        if not exists:
            break
        code = generate_access_code(plan_key)

    # Store code in DB
    with get_db() as conn:
        conn.execute(
            """INSERT INTO codes (code, plan, email, created_at, stripe_session_id, is_subscription, subscription_end, max_users)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (code, plan_name, customer_email, datetime.now().isoformat(), session_id, int(is_subscription))
        )
        conn.execute(
            """INSERT INTO payments (stripe_session_id, customer_email, plan, amount_total, currency, created_at, code_sent)
               VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (session_id, customer_email, plan_name, amount_total, currency, datetime.now().isoformat(), code)
        )
        conn.commit()

    # Send email
    email_sent = send_access_email(customer_email, code, plan_name, lang)
    print(f"[PAYMENT OK] {customer_email} bought {plan_name} → code {code} → email sent: {email_sent}", flush=True)


# ─────────────────────────────────────────────────────────────
# SUCCESS PAGE
# ─────────────────────────────────────────────────────────────

@app.route('/success')
def success():
    try:
        return render_template('success.html')
    except Exception:
        # Fallback inline if template missing
        return """<!DOCTYPE html><html><head><meta charset="UTF-8">
<title>Juriva — Betaling geslaagd</title>
<style>body{font-family:Arial,sans-serif;background:#080808;color:#f0ede8;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0}
.box{max-width:480px;text-align:center;padding:40px 32px;background:#0f0f0f;border:1px solid rgba(255,255,255,0.07);border-radius:12px}
h1{font-size:28px;margin-bottom:12px}p{color:rgba(240,237,232,0.5);line-height:1.7}
.code-note{background:#1a1815;border:1px solid #c8a951;border-radius:8px;padding:16px;margin:24px 0;color:#c8a951;font-size:14px}
a{color:#c8a951;text-decoration:none}</style></head>
<body><div class="box">
<div style="border-top:2px solid #c8a951;width:32px;margin:0 auto 24px"></div>
<h1>Betaling geslaagd</h1>
<div class="code-note">Uw toegangscode is verstuurd naar uw e-mailadres.<br>Controleer ook uw spamfolder.</div>
<p>Ga naar <a href="/">juriva.nl</a> en voer uw code in om direct te beginnen.</p>
<p style="margin-top:24px;font-size:13px">Vragen? <a href="mailto:info@juriva.nl">info@juriva.nl</a></p>
</div></body></html>"""


@app.route('/admin/resend-email', methods=['POST'])
def admin_resend_email():
    """Manually resend access email for a session or email address."""
    data = request.get_json() or {}
    secret = data.get('secret', '')
    if secret != os.environ.get('SECRET_KEY', ''):
        return jsonify({'error': 'Unauthorized'}), 401

    email = data.get('email', '')
    lang = data.get('lang', 'nl')

    with get_db() as conn:
        row = conn.execute(
            "SELECT code, plan FROM codes WHERE email = ? ORDER BY created_at DESC LIMIT 1",
            (email,)
        ).fetchone()

    if not row:
        return jsonify({'error': f'No code found for {email}'}), 404

    code, plan = row['code'], row['plan']
    success = send_access_email(email, code, plan, lang)
    return jsonify({'sent': success, 'code': code, 'plan': plan, 'to': email})


@app.route('/admin/resend/<secret>/<email>')
def admin_resend_get(secret, email):
    """Browser-accessible resend endpoint."""
    if secret != (os.environ.get('SECRET_KEY') or 'REDACTED-ADMIN-SECRET'):
        return "Unauthorized", 401
    with get_db() as conn:
        row = conn.execute(
            "SELECT code, plan FROM codes WHERE email = ? ORDER BY created_at DESC LIMIT 1",
            (email,)
        ).fetchone()
    if not row:
        return f"No code found for {email}", 404
    code, plan = row['code'], row['plan']
    success = send_access_email(email, code, plan, 'nl')
    if success:
        return f"<h2>✓ Email sent!</h2><p>Code <strong>{code}</strong> sent to {email}</p>"
    else:
        return f"<h2>✗ Email failed</h2><p>Check Railway logs for details.</p>", 500

@app.route('/admin/clear-session/<secret>/<session_id>')
def admin_clear_session(secret, session_id):
    """Clear a processed session so webhook can reprocess it."""
    if secret != (os.environ.get('SECRET_KEY') or 'REDACTED-ADMIN-SECRET'):
        return "Unauthorized", 401
    with get_db() as conn:
        conn.execute("DELETE FROM payments WHERE stripe_session_id = ?", (session_id,))
        conn.commit()
    return f"<h2>✓ Session cleared</h2><p>Resend the webhook from Stripe now.</p>"


@app.route('/admin/insert-code/<secret>/<code>/<plan>')
@app.route('/admin/insert-code/<secret>/<code>/<plan>/<int:days>')
def admin_insert_code(secret, code, plan, days=31):
    """Insert a code directly into the DB. Optional days parameter overrides default expiry."""
    if secret != (os.environ.get('SECRET_KEY') or 'REDACTED-ADMIN-SECRET'):
        return "Unauthorized", 401
    is_sub = 0 if plan == 'Per Analyse' else 1
    with get_db() as conn:
        sub_end = None
        if plan in ('ZZP', 'Professioneel', 'Kantoor'):
            sub_end = (datetime.now() + timedelta(days=days)).isoformat()
        max_u = 10 if plan == 'Kantoor' else (3 if plan == 'Professioneel' else 1)
        conn.execute(
            "INSERT OR REPLACE INTO codes (code, plan, email, created_at, is_subscription, analyse_count, subscription_end, max_users) VALUES (?, ?, ?, ?, ?, 0, ?, ?)",
            (code, plan, 'test@juriva.nl', datetime.now().isoformat(), is_sub, sub_end, max_u)
        )
        conn.commit()
    return f"<h2>✓ Code inserted</h2><p>Code: <strong>{code}</strong><br>Plan: <strong>{plan}</strong><br>Valid for: <strong>{days} days</strong><br>Expires: <strong>{sub_end}</strong></p>"


@app.route('/admin/insert-test-code/<secret>/<code>/<plan>')
def admin_insert_test_code(secret, code, plan):
    """Insert a test code with 2-minute expiry for testing."""
    if secret != (os.environ.get('SECRET_KEY') or 'REDACTED-ADMIN-SECRET'):
        return "Unauthorized", 401
    is_sub = 0 if plan == 'Per Analyse' else 1
    sub_end = (datetime.now() + timedelta(minutes=2)).isoformat() if plan not in ('Per Analyse', 'Pilot') else None
    max_u = 10 if plan == 'Kantoor' else (3 if plan == 'Professioneel' else 1)
    with get_db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO codes (code, plan, email, created_at, is_subscription, analyse_count, subscription_end, max_users) VALUES (?, ?, ?, ?, ?, 0, ?, ?)",
            (code, plan, 'test@juriva.nl', datetime.now().isoformat(), is_sub, sub_end, max_u)
        )
        conn.commit()
    return f"<h2>✓ Test code inserted (2-min expiry)</h2><p>Code: <strong>{code}</strong><br>Plan: <strong>{plan}</strong><br>Expires: <strong>{sub_end}</strong></p>"


@app.route('/admin/clear-test-codes/<secret>')
def admin_clear_test_codes(secret):
    """Delete all test codes (email = test@juriva.nl)."""
    if secret != (os.environ.get('SECRET_KEY') or 'REDACTED-ADMIN-SECRET'):
        return "Unauthorized", 401
    with get_db() as conn:
        conn.execute("DELETE FROM codes WHERE email = 'test@juriva.nl'")
        conn.commit()
    return "<h2>✓ All test codes cleared</h2>"


if __name__ == '__main__':
    app.run(debug=True, port=5000)
